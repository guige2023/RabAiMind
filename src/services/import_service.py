"""
Content Import Service
Supports: PDF, DOCX, URL extraction → convert to PPT outline
"""

import io
import logging
from typing import Any

import httpx
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class ImportService:
    """Extract content from various sources and convert to PPT outline format"""

    async def import_pdf(self, file_content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from PDF and convert to outline"""
        try:
            from pypdf import PdfReader
        except ImportError:
            # fallback - try PyPDF2
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                full_text = "\n".join(text_parts)
            except ImportError:
                return {"success": False, "error": "请安装 pypdf: pip install pypdf"}
        else:
            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            full_text = "\n".join(text_parts)

        outline = self._text_to_outline(full_text, source=filename)
        return {
            "success": True,
            "source": "pdf",
            "filename": filename,
            "outline": outline,
            "page_count": len(reader.pages) if 'reader' in dir() else 0
        }

    async def import_docx(self, file_content: bytes, filename: str) -> dict[str, Any]:
        """Extract text from DOCX and convert to outline"""
        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "请安装 python-docx: pip install python-docx"}

        doc = Document(io.BytesIO(file_content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract table content
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        tables_text.append(text)

        full_text = "\n".join(paragraphs + tables_text)
        outline = self._text_to_outline(full_text, source=filename)

        return {
            "success": True,
            "source": "docx",
            "filename": filename,
            "outline": outline,
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables)
        }

    async def import_url(self, url: str) -> dict[str, Any]:
        """Extract content from URL and convert to outline"""
        try:
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                response = await client.get(url)
                response.raise_for_status()

            content_type = response.headers.get("content-type", "")
            if "html" not in content_type and "text" not in content_type:
                return {"success": False, "error": f"URL 内容不是 HTML: {content_type}"}

            html = response.text
            soup = BeautifulSoup(html, "html.parser")

            # Remove script and style elements
            for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
                tag.decompose()

            # Try to get the main content
            article = soup.find("article") or soup.find("main") or soup.find("div", class_="content") or soup

            # Extract title
            title = soup.title.string if soup.title else ""
            if not title:
                h1 = soup.find("h1")
                title = h1.string if h1 else "网页内容"

            # Extract headings and paragraphs
            content_blocks = []

            # Get all headings and their following content
            for tag in article.find_all(["h1", "h2", "h3", "h4", "p", "li"]):
                text = tag.get_text(strip=True)
                if text and len(text) > 10:
                    tag_name = tag.name
                    if tag_name.startswith("h"):
                        content_blocks.append({"type": "heading", "level": int(tag_name[1]), "text": text})
                    elif tag_name == "p":
                        content_blocks.append({"type": "paragraph", "text": text})
                    elif tag_name == "li":
                        content_blocks.append({"type": "list_item", "text": text})

            full_text = "\n".join([b.get("text", "") for b in content_blocks])
            outline = self._text_to_outline(full_text, source=url)

            # Override title with actual page title
            if title:
                outline["title"] = title

            return {
                "success": True,
                "source": "url",
                "url": url,
                "title": title,
                "outline": outline,
                "block_count": len(content_blocks)
            }

        except httpx.TimeoutException:
            return {"success": False, "error": "URL 请求超时，请检查链接是否有效"}
        except httpx.HTTPError as e:
            return {"success": False, "error": f"URL 请求失败: {str(e)}"}
        except Exception as e:
            logger.exception("URL import failed")
            return {"success": False, "error": f"解析失败: {str(e)}"}

    def _text_to_outline(self, text: str, source: str = "") -> dict[str, Any]:
        """Convert extracted text to PPT outline format"""

        # Clean up text
        lines = [line.strip() for line in text.split("\n") if line.strip()]

        # Split into sections by double newlines or major headings
        sections = []
        current_section = {"title": "", "content": []}

        for line in lines:
            # Detect heading-like lines (short, no period ending, or all caps)
            is_heading = (
                len(line) < 80 and
                not line.endswith(".") and
                (line.isupper() or line[0].isupper() and len(line) < 60)
            )

            if is_heading and current_section["content"]:
                # Save current section
                if current_section["title"] or current_section["content"]:
                    sections.append(current_section)
                current_section = {"title": line, "content": []}
            else:
                if current_section["title"] or current_section["content"]:
                    current_section["content"].append(line)
                elif line:
                    # First content without heading - use as intro
                    current_section["title"] = line[:60] if len(line) > 60 else line

        if current_section["title"] or current_section["content"]:
            sections.append(current_section)

        # Build slides
        slides = []

        # Title slide
        title_text = source or "导入内容"
        slides.append({
            "title": title_text,
            "content": "",
            "layout": "title",
            "slide_type": "title"
        })

        # Table of contents if many sections
        if len(sections) > 3:
            toc_content = "\n".join([s["title"] for s in sections[:10] if s["title"]])
            slides.append({
                "title": "目录",
                "content": toc_content,
                "layout": "center",
                "slide_type": "toc"
            })

        # Content slides
        for section in sections[:20]:  # Max 20 content sections
            if not section["title"] and not section["content"]:
                continue

            title = section["title"] or "内容"
            content = "\n".join(section["content"]) if section["content"] else ""

            # Truncate very long content
            if len(content) > 800:
                content = content[:800] + "..."

            slides.append({
                "title": title,
                "content": content,
                "layout": "content",
                "slide_type": "content"
            })

        # Ensure at least one content slide
        if len(slides) == 1:
            slides.append({
                "title": "主要内容",
                "content": "\n".join(lines[:10]) if lines else "暂无详细内容",
                "layout": "content",
                "slide_type": "content"
            })

        return {
            "title": title_text,
            "slides": slides,
            "scene": "general",
            "style": "professional"
        }


    async def import_google_slides(
        self,
        presentation_url: str,
        access_token: str | None = None
    ) -> dict[str, Any]:
        """
        Import content from Google Slides.
        With access_token: Uses Google Slides API to extract slide content.
        Without token: Extracts from publicly shared URL (limited).
        """
        import re

        # Extract presentation ID from URL
        presentation_id = None
        if "presentation/d/" in presentation_url:
            match = re.search(r"presentation/d/([a-zA-Z0-9-_]+)", presentation_url)
            if match:
                presentation_id = match.group(1)
        elif re.match(r"^[a-zA-Z0-9-_]{40,}$", presentation_url):
            presentation_id = presentation_url

        if not presentation_id:
            return {
                "success": False,
                "error": "无法解析 Google Slides URL，请提供有效的分享链接",
                "guide": {
                    "step1": "打开 Google Slides 演示文稿",
                    "step2": "点击 '文件' → '共享' → '发布到网络'",
                    "step3": "复制链接并粘贴到此处"
                }
            }

        if access_token:
            # Use Google Slides API
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    headers = {
                        "Authorization": f"Bearer {access_token}",
                        "Content-Type": "application/json"
                    }
                    # Get presentation metadata
                    resp = await client.get(
                        f"https://slides.googleapis.com/v1/presentations/{presentation_id}",
                        headers=headers
                    )
                    if resp.status_code == 401:
                        return {"success": False, "error": "Google token 无效或已过期"}
                    if resp.status_code == 404:
                        return {"success": False, "error": "演示文稿不存在或无权访问"}
                    resp.raise_for_status()
                    presentation = resp.json()

                    # Get slide content
                    slides_resp = await client.get(
                        f"https://slides.googleapis.com/v1/presentations/{presentation_id}/slides?pageIds=false",
                        headers=headers
                    )
                    slides_data = slides_resp.json() if slides_resp.status_code == 200 else {}

                    title = presentation.get("title", "Google Slides 导入")
                    slides_content = []

                    for i, slide in enumerate(slides_data.get("slides", [])):
                        slide_texts = []
                        for shape in slide.get("pageElements", []):
                            text = shape.get("shape", {}).get("text", {})
                            text_runs = text.get("textRuns", [])
                            for run in text_runs:
                                content = run.get("content", "").strip()
                                if content:
                                    slide_texts.append(content)

                        slide_text = " ".join(slide_texts)
                        slides_content.append({
                            "title": f"第{i+1}页",
                            "content": slide_text,
                            "slide_type": "content"
                        })

                    if not slides_content:
                        return {
                            "success": False,
                            "error": "未能从 Google Slides 提取到文本内容",
                            "hint": "请确保演示文稿包含文本内容（非仅图片）"
                        }

                    # Build outline
                    outline_slides = [{"title": title, "content": "", "layout": "title", "slide_type": "title"}]
                    for sc in slides_content[:30]:
                        outline_slides.append({
                            "title": sc["title"],
                            "content": sc["content"],
                            "layout": "content",
                            "slide_type": "content"
                        })

                    return {
                        "success": True,
                        "source": "google-slides",
                        "presentation_id": presentation_id,
                        "title": title,
                        "outline": {
                            "title": title,
                            "slides": outline_slides,
                            "scene": "general",
                            "style": "professional"
                        },
                        "slides_count": len(slides_content)
                    }
            except httpx.HTTPError as e:
                logger.error(f"Google Slides API error: {e}")
                return {"success": False, "error": f"Google Slides API 请求失败: {str(e)}"}
            except Exception:
                logger.exception("Google Slides import error")
                return {"success": False, "error": "导入失败"}
        else:
            # No token - try public export endpoint
            try:
                export_url = f"https://docs.google.com/presentation/d/{presentation_id}/export/pdf"
                return {
                    "success": False,
                    "error": "Google Slides 导入需要授权",
                    "guide": {
                        "step1": "此功能需要 Google OAuth 授权",
                        "step2": "请先在设置中连接 Google 账号",
                        "step3": "授权后将自动提取演示文稿内容"
                    },
                    "presentation_id": presentation_id,
                    "requires_auth": True
                }
            except Exception:
                return {"success": False, "error": "Google Slides URL 无效"}

    async def import_pinterest(
        self,
        board_url: str,
        access_token: str | None = None
    ) -> dict[str, Any]:
        """
        Import content from Pinterest board.
        With access_token: Uses Pinterest API to fetch board pins.
        Without token: Scrapes publicly available board content.
        """
        import re
        from urllib.parse import urlparse

        # Extract board ID or username from URL
        board_id = None
        parsed = urlparse(board_url)

        if "pinterest.com" in parsed.netloc:
            path_parts = [p for p in parsed.path.split("/") if p]
            if len(path_parts) >= 2 and path_parts[0] in ("board", "person", "username"):
                board_id = path_parts[1] if path_parts[0] == "board" else None
            elif len(path_parts) >= 1:
                # Could be a board URL like /board/abc123/
                for i, part in enumerate(path_parts):
                    if part == "board" and i + 1 < len(path_parts):
                        board_id = path_parts[i + 1]
                        break

        # Also accept direct board ID
        if not board_id and re.match(r"^[a-zA-Z0-9_-]{10,}$", board_url.strip()):
            board_id = board_url.strip()

        if not board_id:
            return {
                "success": False,
                "error": "无法解析 Pinterest board URL",
                "guide": {
                    "step1": "打开 Pinterest board 页面",
                    "step2": "复制浏览器地址栏的 URL",
                    "step3": "粘贴到此处"
                }
            }

        if access_token:
            # Use Pinterest API
            try:
                async with httpx.AsyncClient(timeout=60.0) as client:
                    headers = {
                        "Authorization": f"Bearer {access_token}",
                        "Content-Type": "application/json"
                    }
                    # Get board info
                    board_resp = await client.get(
                        f"https://api.pinterest.com/v5/boards/{board_id}",
                        headers=headers
                    )
                    if board_resp.status_code == 401:
                        return {"success": False, "error": "Pinterest token 无效或已过期"}
                    if board_resp.status_code == 404:
                        return {"success": False, "error": "Board 不存在或无权访问"}
                    board_resp.raise_for_status()
                    board_data = board_resp.json()

                    board_name = board_data.get("name", "Pinterest Board")

                    # Get pins
                    pins_resp = await client.get(
                        f"https://api.pinterest.com/v5/boards/{board_id}/pins",
                        headers=headers,
                        params={"page_size": 50}
                    )
                    pins_data = pins_resp.json() if pins_resp.status_code == 200 else {}
                    pins = pins_data.get("items", [])

                    slides = [{"title": board_name, "content": "", "layout": "title", "slide_type": "title"}]

                    for i, pin in enumerate(pins[:30]):
                        title = pin.get("title", f"第{i+1}张图")
                        description = pin.get("description", "")
                        link = pin.get("link", "")

                        content_parts = []
                        if description:
                            content_parts.append(description)
                        if link:
                            content_parts.append(f"链接: {link}")

                        slides.append({
                            "title": title or f"第{i+1}张图",
                            "content": "\n".join(content_parts) if content_parts else "",
                            "layout": "center",
                            "slide_type": "content"
                        })

                    return {
                        "success": True,
                        "source": "pinterest",
                        "board_id": board_id,
                        "board_name": board_name,
                        "outline": {
                            "title": board_name,
                            "slides": slides,
                            "scene": "general",
                            "style": "professional"
                        },
                        "pins_count": len(pins)
                    }
            except httpx.HTTPError as e:
                logger.error(f"Pinterest API error: {e}")
                return {"success": False, "error": f"Pinterest API 请求失败: {str(e)}"}
            except Exception:
                logger.exception("Pinterest import error")
                return {"success": False, "error": "导入失败"}
        else:
            # No token - try to scrape public board
            try:
                async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                    response = await client.get(board_url)
                    response.raise_for_status()

                    html = response.text
                    soup = BeautifulSoup(html, "html.parser")

                    # Try to extract board name
                    title_tag = soup.find("title")
                    board_name = title_tag.string.replace(" - Pinterest", "").strip() if title_tag else "Pinterest Board"

                    # Extract pin titles and descriptions from JSON-LD or meta tags
                    slides = [{"title": board_name, "content": "", "layout": "title", "slide_type": "title"}]

                    # Pinterest embeds data in script tags
                    pin_data = []
                    for script in soup.find_all("script"):
                        text = script.string or ""
                        if "pinId" in text or '"description"' in text:
                            pin_data.append(text)

                    if pin_data:
                        import json
                        for script_text in pin_data[:5]:  # Limit processing
                            try:
                                # Try to find JSON data
                                start = script_text.find("{")
                                end = script_text.rfind("}") + 1
                                if start >= 0 and end > start:
                                    data = json.loads(script_text[start:end])
                                    if isinstance(data, dict):
                                        desc = data.get("description", "") or data.get("headline", "")
                                        if desc:
                                            slides.append({
                                                "title": data.get("name", data.get("headline", "图片")),
                                                "content": str(desc),
                                                "layout": "center",
                                                "slide_type": "content"
                                            })
                            except (json.JSONDecodeError, ValueError):
                                pass

                    if len(slides) == 1:
                        return {
                            "success": False,
                            "error": "未能从 Pinterest board 提取到内容",
                            "guide": {
                                "step1": "此功能需要 Pinterest OAuth 授权",
                                "step2": "请先在设置中连接 Pinterest 账号",
                                "step3": "授权后将自动提取 Board 中的图钉内容"
                            },
                            "board_id": board_id,
                            "requires_auth": True
                        }

                    return {
                        "success": True,
                        "source": "pinterest",
                        "board_id": board_id,
                        "board_name": board_name,
                        "outline": {
                            "title": board_name,
                            "slides": slides[:31],  # Max 30 content slides
                            "scene": "general",
                            "style": "professional"
                        }
                    }
            except httpx.HTTPError as e:
                return {
                    "success": False,
                    "error": f"无法访问 Pinterest board: {str(e)}",
                    "guide": {
                        "step1": "请确保 Board 是公开的",
                        "step2": "或者先在设置中授权 Pinterest 账号"
                    }
                }
            except Exception:
                logger.exception("Pinterest scrape error")
                return {"success": False, "error": "解析 Pinterest board 失败"}

    async def import_images(
        self,
        image_data_list: list[tuple] | None = None,
        titles: list[str] | None = None,
        captions: list[str] | None = None,
        layout: str = "center",
        scene: str = "general",
        style: str = "professional"
    ) -> dict[str, Any]:
        """
        Import images (uploaded as bytes) and convert to PPT outline.
        Each image becomes one slide.

        Args:
            image_data_list: List of (filename, bytes) tuples
            titles: Optional list of titles per image
            captions: Optional list of captions per image
        """
        if image_data_list is None:
            image_data_list = []
        if not titles:
            titles = []
        if not captions:
            captions = []

        num_images = len(image_data_list)
        if num_images == 0:
            return {
                "success": False,
                "error": "没有提供图片，请上传至少一张图片"
            }

        # Save images to temp directory for reference
        import os
        import tempfile
        temp_dir = tempfile.mkdtemp(prefix="rabai_import_images_")
        saved_image_paths = []

        for i, (filename, image_bytes) in enumerate(image_data_list):
            safe_name = f"img_{i+1:03d}_{filename}"
            img_path = os.path.join(temp_dir, safe_name)
            try:
                with open(img_path, "wb") as f:
                    f.write(image_bytes)
                saved_image_paths.append(img_path)
            except Exception as e:
                logger.warning(f"Failed to save image {filename}: {e}")

        slides = []
        for i in range(num_images):
            img_title = titles[i] if i < len(titles) and titles[i] else f"图片 {i+1}"
            img_caption = captions[i] if i < len(captions) else ""
            img_path = saved_image_paths[i] if i < len(saved_image_paths) else None

            slides.append({
                "title": img_title,
                "content": img_caption,
                "layout": layout,
                "slide_type": "content",
                "image_path": img_path
            })

        return {
            "success": True,
            "source": "images",
            "count": num_images,
            "layout": layout,
            "images_dir": temp_dir,
            "outline": {
                "title": f"图片集 ({num_images} 张)",
                "slides": slides,
                "scene": scene,
                "style": style
            }
        }


    async def import_notion(
        self,
        page_url: str,
        access_token: str | None = None
    ) -> dict[str, Any]:
        """
        Import content from Notion page.
        Requires Notion integration token (from https://www.notion.so/my-integrations).
        """
        import re

        # Extract page ID from URL
        page_id = None
        # Patterns: notion.so/username/Page-Title-1234567890abcdef...
        match = re.search(r"notion\.so/[\w-]+/([a-f0-9]{32})", page_url)
        if match:
            page_id = match.group(1)
        else:
            # Try direct page ID (32 hex chars)
            match = re.search(r"^([a-f0-9]{32})$", page_url.strip())
            if match:
                page_id = match.group(1)

        if not page_id:
            return {
                "success": False,
                "error": "无法解析 Notion 页面 URL，请确保链接格式正确",
                "guide": {
                    "step1": "打开要导入的 Notion 页面",
                    "step2": "点击右上角 '...' 菜单 → 复制链接",
                    "step3": "确保 Notion 集成已添加到该页面（通过 '...' → 添加连接）",
                }
            }

        if not access_token:
            # Try environment variable
            import os
            access_token = os.environ.get("NOTION_TOKEN")

        if not access_token:
            return {
                "success": False,
                "error": "Notion 导入需要 API Token",
                "guide": {
                    "step1": "前往 https://www.notion.so/my-integrations 创建集成",
                    "step2": "复制 Integration Token（secret_xxx）",
                    "step3": "在 Notion 页面中，点击 '...' → 添加连接 → 选择你的集成",
                    "step4": "重新粘贴 Notion 页面链接并导入",
                },
                "requires_auth": True,
                "page_id": page_id
            }

        try:
            headers = {
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
                "Notion-Version": "2022-06-28"
            }

            async with httpx.AsyncClient(timeout=60.0) as client:
                # Get page info
                page_resp = await client.get(
                    f"https://api.notion.com/v1/pages/{page_id}",
                    headers=headers
                )
                if page_resp.status_code == 404:
                    return {
                        "success": False,
                        "error": "页面不存在或无权访问，请确保已在 Notion 页面中添加该集成"
                    }
                if page_resp.status_code == 401:
                    return {
                        "success": False,
                        "error": "Notion Token 无效或已过期，请检查集成设置"
                    }
                page_resp.raise_for_status()
                page_data = page_resp.json()

                # Get page title from properties
                title = "Notion 内容"
                props = page_data.get("properties", {})
                for prop_name, prop_value in props.items():
                    if prop_value.get("type") == "title":
                        title_texts = prop_value.get("title", [])
                        if title_texts:
                            title = "".join(t.get("plain_text", "") for t in title_texts)
                            break

                # Get page blocks (content) - use GET not POST
                blocks_resp = await client.get(
                    f"https://api.notion.com/v1/blocks/{page_id}/children",
                    headers=headers,
                    params={"page_size": 100}
                )
                blocks_data = blocks_resp.json() if blocks_resp.status_code == 200 else {}
                blocks = blocks_data.get("results", [])

                # Extract text from blocks
                content_blocks = []
                for block in blocks:
                    block_type = block.get("type", "")
                    block_content = block.get(block_type, {})

                    # Get text from various block types
                    text_runs = []
                    if "rich_text" in block_content:
                        text_runs = block_content["rich_text"]
                    elif block_type == "bulleted_list_item":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "numbered_list_item":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "to_do":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "heading_1":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "heading_2":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "heading_3":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "paragraph":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "quote":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "callout":
                        text_runs = block_content.get("rich_text", [])
                    elif block_type == "code":
                        text_runs = block_content.get("rich_text", [])

                    text = "".join(run.get("plain_text", "") for run in text_runs)
                    if text and len(text) > 3:
                        content_blocks.append({"type": block_type, "text": text})

                # Build full text for outline conversion
                full_text = "\n".join([b["text"] for b in content_blocks])
                outline = self._text_to_outline(full_text, source=title)
                outline["title"] = title

                return {
                    "success": True,
                    "source": "notion",
                    "page_id": page_id,
                    "title": title,
                    "outline": outline,
                    "block_count": len(content_blocks)
                }

        except httpx.HTTPError as e:
            logger.error(f"Notion API error: {e}")
            return {"success": False, "error": f"Notion API 请求失败: {str(e)}"}
        except Exception:
            logger.exception("Notion import error")
            return {"success": False, "error": "导入 Notion 页面失败"}

    async def import_google_docs(
        self,
        doc_url: str,
        access_token: str | None = None
    ) -> dict[str, Any]:
        """
        Import content from Google Docs (not Slides).
        With access_token: Uses Google Docs API to extract content.
        Without token: Returns helpful error with steps.
        """
        import re

        # Extract document ID from URL
        doc_id = None
        # https://docs.google.com/document/d/DOC_ID/edit
        match = re.search(r"/document/d/([a-zA-Z0-9-_]+)", doc_url)
        if match:
            doc_id = match.group(1)
        elif re.match(r"^[a-zA-Z0-9-_]{30,}$", doc_url.strip()):
            doc_id = doc_url.strip()

        if not doc_id:
            return {
                "success": False,
                "error": "无法解析 Google Docs URL，请确保是 docs.google.com/document/ 链接",
                "guide": {
                    "step1": "打开 Google Docs 文档",
                    "step2": "点击 '文件' → '共享' → '发布到网络'",
                    "step3": "或者复制浏览器地址栏中的文档链接",
                }
            }

        if not access_token:
            return {
                "success": False,
                "error": "Google Docs 导入需要授权",
                "guide": {
                    "step1": "此功能需要 Google OAuth 授权",
                    "step2": "请先在设置中连接 Google 账号",
                    "step3": "授权后将自动提取文档内容"
                },
                "requires_auth": True,
                "doc_id": doc_id
            }

        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                headers = {
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json"
                }

                # Get document metadata
                doc_resp = await client.get(
                    f"https://docs.googleapis.com/v1/documents/{doc_id}",
                    headers=headers
                )
                if doc_resp.status_code == 404:
                    return {"success": False, "error": "文档不存在或无权访问"}
                if doc_resp.status_code == 401:
                    return {"success": False, "error": "Google token 无效或已过期"}
                doc_resp.raise_for_status()
                doc_data = doc_resp.json()

                title = doc_data.get("title", "Google Docs 文档")

                # Get document body content
                body = doc_data.get("body", {})
                content_blocks = []

                def extract_text_from_content(content_item):
                    """Recursively extract text from content elements."""
                    results = []
                    if "textRun" in content_item:
                        text = content_item["textRun"].get("content", "")
                        if text.strip():
                            results.append(text.strip())
                    if "paragraph" in content_item:
                        para = content_item["paragraph"]
                        elements = para.get("elements", [])
                        for elem in elements:
                            if "textRun" in elem:
                                text = elem["textRun"].get("content", "")
                                if text.strip():
                                    results.append(text.strip())
                    if "table" in content_item:
                        table = content_item["table"]
                        for row in table.get("tableRows", []):
                            for cell in row.get("tableCells", []):
                                for content in cell.get("content", []):
                                    results.extend(extract_text_from_content(content))
                    if "sectionBreak" in content_item:
                        results.append("===")  # Section divider
                    return results

                # Process body content
                body_content = body.get("content", [])
                for item in body_content:
                    if "paragraph" in item:
                        para_texts = extract_text_from_content(item)
                        if para_texts:
                            text = " ".join(para_texts)
                            if text.strip():
                                # Determine block type by paragraph style
                                para_props = item["paragraph"].get("paragraphStyle", {})
                                named_style = para_props.get("namedStyleType", "NORMAL_TEXT")

                                if "HEADING" in named_style or "TITLE" in named_style:
                                    content_blocks.append({"type": "heading", "text": text})
                                else:
                                    content_blocks.append({"type": "paragraph", "text": text})
                    elif "table" in item:
                        table_texts = []
                        for row in item["table"].get("tableRows", []):
                            for cell in row.get("tableCells", []):
                                cell_texts = []
                                for content in cell.get("content", []):
                                    cell_texts.extend(extract_text_from_content(content))
                                if cell_texts:
                                    table_texts.append(" | ".join(cell_texts))
                        if table_texts:
                            content_blocks.append({"type": "table", "text": "\n".join(table_texts)})

                if not content_blocks:
                    return {
                        "success": False,
                        "error": "未能从 Google Docs 提取到文本内容"
                    }

                # Build text for outline conversion
                full_text = "\n".join([b["text"] for b in content_blocks if b["text"].strip()])
                outline = self._text_to_outline(full_text, source=title)
                outline["title"] = title

                return {
                    "success": True,
                    "source": "google-docs",
                    "doc_id": doc_id,
                    "title": title,
                    "outline": outline,
                    "block_count": len(content_blocks)
                }

        except httpx.HTTPError as e:
            logger.error(f"Google Docs API error: {e}")
            return {"success": False, "error": f"Google Docs API 请求失败: {str(e)}"}
        except Exception:
            logger.exception("Google Docs import error")
            return {"success": False, "error": "导入 Google Docs 失败"}


def get_import_service() -> ImportService:
    return ImportService()
